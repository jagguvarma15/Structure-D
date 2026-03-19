"""Shared fixtures for the Structure-D test suite."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Type

import pytest
from pydantic import BaseModel

from structure_d.inference.providers import BaseLLMProvider, ProviderResult
from structure_d.schemas.base import (
    ChunkMetadata,
    DocumentFormat,
    DocumentMetadata,
    ExtractionResult,
    ParsedDocument,
    TaskType,
    TextChunk,
)
from structure_d.schemas.generic import KeyValueExtraction, KeyValuePair


# ── Fixture files ────────────────────────────────────────────────────────────


@pytest.fixture()
def sample_text_file(tmp_path: Path) -> Path:
    """A tiny .txt file for ingestion tests."""
    f = tmp_path / "sample.txt"
    f.write_text("Name: Alice\nAge: 30\nEmail: alice@example.com")
    return f


@pytest.fixture()
def sample_md_file(tmp_path: Path) -> Path:
    """A Markdown file with headings for chunker tests."""
    f = tmp_path / "sample.md"
    f.write_text("# Title\n\nFirst paragraph.\n\n## Section\n\nSecond paragraph.")
    return f


@pytest.fixture()
def sample_html_file(tmp_path: Path) -> Path:
    """A minimal HTML file for HTML parser tests."""
    f = tmp_path / "sample.html"
    f.write_text(
        "<html><head><title>Test</title></head>"
        "<body><p>Hello World</p><script>var x=1;</script></body></html>"
    )
    return f


@pytest.fixture()
def sample_email_file(tmp_path: Path) -> Path:
    """A minimal .eml file for email parser tests."""
    f = tmp_path / "sample.eml"
    f.write_text(
        "From: sender@example.com\n"
        "To: receiver@example.com\n"
        "Subject: Test Email\n"
        "Date: Thu, 12 Mar 2026 10:00:00 +0000\n"
        "Content-Type: text/plain; charset=utf-8\n"
        "\n"
        "This is the email body.\n"
    )
    return f


@pytest.fixture()
def sample_csv_file(tmp_path: Path) -> Path:
    """A tiny CSV file for plain-text parser tests."""
    f = tmp_path / "sample.csv"
    f.write_text("name,age\nAlice,30\nBob,25\n")
    return f


@pytest.fixture()
def sample_docx_file(tmp_path: Path) -> Path:
    """A .docx file with headings, paragraphs, and a table for parser tests."""
    from docx import Document

    doc = Document()
    doc.add_heading("Sample Report", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("Structure-D extracts structured data from documents.")
    doc.add_paragraph("It supports multiple file formats.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "vendor"
    table.cell(1, 1).text = "Acme Corp"
    table.cell(2, 0).text = "total"
    table.cell(2, 1).text = "1240.00"

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


# ── Schema / result helpers ──────────────────────────────────────────────────


@pytest.fixture()
def kv_extraction() -> KeyValueExtraction:
    """A pre-built KeyValueExtraction instance."""
    return KeyValueExtraction(
        pairs=[
            KeyValuePair(key="name", value="Alice", confidence=0.99),
            KeyValuePair(key="age", value="30", confidence=0.95),
        ]
    )


@pytest.fixture()
def extraction_result() -> ExtractionResult:
    """A valid ExtractionResult for storage / pipeline tests."""
    return ExtractionResult(
        document_id="doc-1",
        chunk_id="chunk-1",
        source_format=DocumentFormat.PLAIN_TEXT,
        task=TaskType.EXTRACTION,
        model_used="test-model",
        raw_output='{"pairs": [{"key": "name", "value": "Alice"}]}',
        structured_output={"pairs": [{"key": "name", "value": "Alice"}]},
        is_valid=True,
        latency_ms=42.0,
        token_usage={"prompt_tokens": 10, "completion_tokens": 20},
    )


@pytest.fixture()
def parsed_document() -> ParsedDocument:
    """A ParsedDocument matching a plain-text ingestion."""
    return ParsedDocument(
        metadata=DocumentMetadata(
            document_id="doc-1",
            filename="sample.txt",
            source="/tmp/sample.txt",
            file_extension=".txt",
            format=DocumentFormat.PLAIN_TEXT,
            file_size_bytes=45,
        ),
        text="Name: Alice\nAge: 30\nEmail: alice@example.com",
        pages=["Name: Alice\nAge: 30\nEmail: alice@example.com"],
    )


@pytest.fixture()
def text_chunk() -> TextChunk:
    """A single TextChunk for inference / validation tests."""
    return TextChunk(
        text="Name: Alice\nAge: 30",
        metadata=ChunkMetadata(
            chunk_id="chunk-1",
            document_id="doc-1",
            source_format=DocumentFormat.PLAIN_TEXT,
            token_count=8,
        ),
    )


# ── Mock provider ────────────────────────────────────────────────────────────


class FakeProvider(BaseLLMProvider):
    """In-memory provider that returns a fixed KeyValueExtraction."""

    def __init__(self, output: BaseModel | None = None, raw: str | None = None) -> None:
        kv = output or KeyValueExtraction(
            pairs=[KeyValuePair(key="name", value="Alice")]
        )
        self._output = kv
        self._raw = raw or json.dumps(kv.model_dump())

    async def generate(
        self,
        prompt: str,
        schema: Type[BaseModel],
        system_prompt: str | None = None,
        temperature: float = 0.0,
        max_tokens: int = 2048,
        model: str | None = None,
        **kwargs: Any,
    ) -> ProviderResult:
        return ProviderResult(
            output=self._output,
            raw_text=self._raw,
            model_used=model or "fake-model",
            token_usage={"prompt_tokens": 10, "completion_tokens": 20},
        )

    async def generate_text(
        self,
        prompt: str,
        system_prompt: str | None = None,
        temperature: float = 0.7,
        max_tokens: int = 2048,
        model: str | None = None,
        **kwargs: Any,
    ) -> str:
        return self._raw


@pytest.fixture()
def fake_provider() -> FakeProvider:
    """A deterministic LLM provider that requires no network."""
    return FakeProvider()
